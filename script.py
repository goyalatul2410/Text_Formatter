from docx import Document
from docx.shared import RGBColor
import re


def textColor(s, clr):
    run = para.add_run(s)
    run.bold = True
    font = run.font
    if clr == "blue":
        font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    if clr == "purple":
        font.color.rgb = RGBColor(0x80, 0x00, 0x80)
    if clr == "green":
        font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    para.add_run(" ")


doc = Document('test.docx')

for para in doc.paragraphs:
    word = para.text.split(" ")
    s1 = "JasperGold"
    s2 = "@"
    x = None
    jasNext = None
    atNext1 = None
    atNext2 = None
    atNext3 = None
    count = 0
    para.clear()

    for s in word:
        if s1 == s:
            x = word.index(s1)
            jasNext = word[x + 1]
        if s2 in s:
            if s2 == s:
                y = word.index(s2)
                atNext1 = word[y + 1]
                atNext2 = word[y + 2]
                atNext3 = word[y + 3]
            else:
                s_split = s.split("@")
                atNext1 = s_split[1]
                y = word.index(s)
                atNext2 = word[y + 1]
                atNext3 = word[y + 2]

    for s in word:

        # Bold
        if jasNext == s and count == 0:
            # print(jasNext)
            para.add_run(jasNext).bold = True
            para.add_run(" ")
            count = 1
        # Bold

        # Blue and Bold
        elif str(atNext1) in s and count == 1:
            if "INTERNAL" in atNext1:
                # print(atNext1)
                para.add_run(s)
                para.add_run(" ")
            elif "@" in s:
                para.add_run("@")
                para.add_run(" ")
                textColor(atNext1, "blue")
                count = 2
            else:
                # print(atNext1)
                textColor(atNext1, "blue")
                count = 2
        elif atNext2 == s and count == 2:
            if "(" in atNext2:
                para.add_run(s)
                para.add_run(" ")
                count = 3
            else:
                textColor(atNext2, "blue")
                count = 3
        elif atNext3 == s and count == 3:
            if "(" in atNext3:
                para.add_run(s)
                para.add_run(" ")
                count = 4
            else:
                if "(" in atNext2:
                    para.add_run(s)
                    para.add_run(" ")
                    count = 4
                else:
                    textColor(atNext3, "blue")
                    count = 4
        # Blue and Bold

        # Purple and Bold
        elif re.search("Synopsys", s, re.IGNORECASE):
            textColor(s, "purple")
        elif re.search("Spyglass", s, re.IGNORECASE):
            textColor(s, "purple")
        elif re.search("VCFormal", s, re.IGNORECASE):
            textColor(s, "purple")
        # Purple and Bold

        #  Green and Bold
        elif re.search("mentor", s, re.IGNORECASE):
            textColor(s, "green")
        elif re.search("0-in", s, re.IGNORECASE):
            textColor(s, "green")
        elif re.search("questa", s, re.IGNORECASE):
            textColor(s, "green")
        elif re.search("formal", s, re.IGNORECASE):
            q = word.index(s)
            if word[q - 1] == 'Questa':
                textColor(word[q], "green")
        # Green and Bold

        # Plain Text
        else:
            para.add_run(s)
            para.add_run(" ")
        # Plain Text

doc.save('demo.docx')